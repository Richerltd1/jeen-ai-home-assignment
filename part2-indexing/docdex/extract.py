"""Text extraction from PDF and DOCX files.

Extraction is deliberately separated from cleaning: `_read_pdf` / `_read_docx`
know about file formats, `clean_text` knows about whitespace, and neither knows
about chunking. That keeps a new format (say .txt) to a single new reader.
"""

from __future__ import annotations

import re
from pathlib import Path

from .errors import (
    FileNotFoundErrorDocdex,
    NoTextExtractedError,
    UnsupportedFileTypeError,
)

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


def extract_text(path: Path) -> str:
    """Extract clean plain text from a PDF or DOCX file.

    Args:
        path: Path to the source document.

    Returns:
        Cleaned document text.

    Raises:
        FileNotFoundErrorDocdex: the path does not exist or is not a file.
        UnsupportedFileTypeError: the extension is not .pdf or .docx.
        NoTextExtractedError: the file parsed but yielded no usable text.
    """
    if not path.exists() or not path.is_file():
        raise FileNotFoundErrorDocdex(f"File not found: {path}")

    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise UnsupportedFileTypeError(
            f"Unsupported file type '{suffix or '(none)'}'. Supported types: {supported}."
        )

    is_pdf = suffix == ".pdf"
    raw = _read_pdf(path) if is_pdf else _read_docx(path)
    text = clean_text(raw)

    # PDFs carry hard line wraps but no paragraph markers, so paragraph structure
    # has to be reconstructed. DOCX already knows its own paragraphs.
    if is_pdf:
        text = reflow_paragraphs(text)

    if not text:
        raise NoTextExtractedError(
            f"No extractable text found in {path.name}. "
            "If this is a scanned document it would need OCR, which is not supported."
        )
    return text


def _read_pdf(path: Path) -> str:
    """Read every page of a PDF, tolerating individual unreadable pages."""
    from pypdf import PdfReader
    from pypdf.errors import PdfReadError

    try:
        reader = PdfReader(str(path))
    except (PdfReadError, OSError, ValueError) as exc:
        raise NoTextExtractedError(f"Could not parse PDF {path.name}: {exc}") from exc

    pages: list[str] = []
    for page in reader.pages:
        try:
            pages.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001 - one broken page must not kill the run
            continue

    # Blank line between pages so paragraph splitting still sees a boundary.
    return "\n\n".join(pages)


def _read_docx(path: Path) -> str:
    """Read paragraphs and table cells from a DOCX file."""
    import docx
    from docx.opc.exceptions import PackageNotFoundError

    try:
        document = docx.Document(str(path))
    except (PackageNotFoundError, OSError, ValueError) as exc:
        raise NoTextExtractedError(f"Could not parse DOCX {path.name}: {exc}") from exc

    blocks = [p.text for p in document.paragraphs]

    # Table text is real content and is easy to lose -- pull it in explicitly.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))

    return "\n\n".join(blocks)


def reflow_paragraphs(text: str) -> str:
    """Rebuild paragraph structure from hard-wrapped PDF lines.

    A PDF stores each visual *line*, not each paragraph, so extracted text has a
    newline every ~90 characters and no blank line between paragraphs. Without
    reconstruction the paragraph strategy sees one enormous block and degenerates
    into the fixed-size strategy, which would make two of the three required
    strategies identical.

    The heuristic exploits how wrapped text looks: every line in a paragraph is
    filled to roughly the column width *except the last one*, which is ragged. So
    a short line ends a paragraph. Numbered or unpunctuated short lines are
    treated as headings and stand alone.
    """
    import statistics

    lines = [line.strip() for line in text.split("\n")]

    # Median over substantial lines only, so headings and stubs do not drag the
    # estimate of the column width down.
    body_lengths = [len(line) for line in lines if len(line) > 20]
    if not body_lengths:
        return text
    ragged_threshold = statistics.median(body_lengths) * 0.75

    paragraphs: list[str] = []
    buffer: list[str] = []

    def flush() -> None:
        if buffer:
            paragraphs.append(" ".join(buffer))
            buffer.clear()

    for line in lines:
        if not line:
            flush()
            continue

        if _is_heading(line):
            flush()
            paragraphs.append(line)
            continue

        buffer.append(line)

        # A line that does not fill the column is the last line of its paragraph.
        if len(line) < ragged_threshold:
            flush()

    flush()
    return "\n\n".join(paragraphs)


def _is_heading(line: str) -> bool:
    """Detect a standalone heading line: numbered, or short and unpunctuated."""
    if re.match(r"^\d+(\.\d+)*\.?\s+\S", line):
        return True
    return len(line) < 60 and not line.endswith((".", ",", ";", ":", "!", "?"))


def clean_text(raw: str) -> str:
    """Normalise whitespace while preserving paragraph boundaries.

    Paragraph breaks survive because the paragraph chunking strategy depends on
    them; everything else collapses to single spaces.
    """
    if not raw:
        return ""

    text = raw.replace("\r\n", "\n").replace("\r", "\n")

    # De-hyphenate words broken across a line break ("inter-\nnational").
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)

    # Drop control characters that survive some PDF encoders.
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)

    # Collapse runs of spaces/tabs, then normalise blank-line runs to exactly two
    # newlines so paragraph detection has a single unambiguous separator.
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r" *\n *", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()
